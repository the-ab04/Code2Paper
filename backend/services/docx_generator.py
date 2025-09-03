from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def render_docx(sections, output_path, title=None, author=None, style=None):
    """
    Generates a DOCX research paper from provided sections.
    :param sections: dict with keys like ['abstract', 'introduction', 'methods', 'results', 'conclusion', 'references']
    :param output_path: path to save the generated DOCX file
    :param title: Paper title
    :param author: Author name
    :param style: Optional style (APA, IEEE, etc.)
    """

    doc = Document()

    # ✅ Title Page
    if title:
        title_paragraph = doc.add_paragraph(title)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.style = doc.styles['Title']

    if author:
        author_paragraph = doc.add_paragraph(author)
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_paragraph.style = doc.styles['Normal']
        doc.add_paragraph()  # Blank line

    # ✅ Sections in order
    section_order = ["abstract", "introduction", "methods", "results", "conclusion"]

    for sec in section_order:
        if sec in sections and sections[sec]:
            # Add heading
            doc.add_heading(sec.capitalize(), level=1)
            # Add content
            doc.add_paragraph(sections[sec])

    # ✅ References
    if "references" in sections and sections["references"]:
        doc.add_heading("References", level=1)
        if isinstance(sections["references"], list):
            for ref in sections["references"]:
                doc.add_paragraph(ref, style="Normal")
        else:
            doc.add_paragraph(sections["references"], style="Normal")

    # ✅ Apply basic font size (11pt)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

    # Save DOCX
    doc.save(output_path)
