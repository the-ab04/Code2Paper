# backend/services/file_generator.py

import io
import os
import re
import json
from typing import Dict, Union, List, Optional, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Optional Pillow import for robust image handling (size detection/resizing)
try:
    from PIL import Image
except Exception:
    Image = None

# Tunables (environment override)
DEFAULT_MAX_FIGURES = int(os.getenv("MAX_FIGURES", "6"))
MIN_IMAGE_BYTES = int(os.getenv("MIN_IMAGE_BYTES", str(8 * 1024)))  # 8KB default
# Minimum characters for a paragraph; if a paragraph smaller than this appears,
# we'll try to merge it with the following paragraph to avoid very short paras.
MIN_PARAGRAPH_CHARS = int(os.getenv("MIN_PARAGRAPH_CHARS", "120"))


def render_docx(
    sections: Dict[str, str],
    out_file: Union[str, io.BytesIO],
    title: str = "Generated Paper",
    images: Optional[List[Dict[str, Any]]] = None,
    max_image_width_inch: float = 6.0,
) -> None:
    """
    Render a DOCX file from structured paper sections.

    Parameters
    ----------
    sections : dict
        mapping canonical section keys -> section text. May include:
          - "images": optional list of image dicts (fallback)
          - "figure_captions": optional JSON array of captions (string) produced by LLM
    out_file : str | io.BytesIO
        path or BytesIO to write the .docx to
    title : str
        document title
    images : optional list
        list of image dicts overriding sections["images"]. Each dict may contain:
           see original format in the repo.
    max_image_width_inch : float
        maximum width to insert images in inches
    """
    doc = Document()

    # === Title ===
    if title and title.strip():
        h = doc.add_heading(title.strip(), level=0)
        # center the title for aesthetics where possible
        try:
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception:
            pass

    # canonical order and mapping for nice headings
    ordered_sections = [
        "abstract",
        "introduction",
        "literature_review",
        "methods",
        "experiments",
        "results",
        "conclusion",
    ]

    def pretty_heading(key: str) -> str:
        return key.replace("_", " ").strip().title()

    # ---------- Paragraph splitting helpers ----------
    def _sentence_split(text: str) -> List[str]:
        """
        Split text into sentences in a tolerant way.
        Uses a simple regex to split at sentence boundaries while keeping abbreviations fairly safe.
        """
        if not text:
            return []
        # Normalize whitespace
        t = re.sub(r"\s+", " ", text.strip())
        # Use lookbehind to split on sentence enders followed by space and capital letter or digit.
        parts = re.split(r'(?<=[\.!?])\s+(?=[A-Z0-9"\'\u2018\u201C])', t)
        return [p.strip() for p in parts if p.strip()]

    def _chunk_sentences_into_paragraphs(sentences: List[str], max_chars: int = 500) -> List[str]:
        """
        Group consecutive sentences into paragraphs where each paragraph is at most max_chars.
        """
        if not sentences:
            return []
        paras = []
        current = []
        current_len = 0
        for s in sentences:
            s_len = len(s)
            # if adding this sentence would exceed the budget, flush current
            if current and (current_len + 1 + s_len) > max_chars:
                paras.append(" ".join(current).strip())
                current = [s]
                current_len = s_len
            else:
                current.append(s)
                current_len += (s_len + (1 if current else 0))
        if current:
            paras.append(" ".join(current).strip())
        return paras

    def _merge_short_paragraphs(paragraphs: List[str], min_chars: int = MIN_PARAGRAPH_CHARS) -> List[str]:
        """
        Merge paragraphs smaller than min_chars with the next paragraph (if available)
        to avoid many very short paragraphs produced by the splitter.
        """
        if not paragraphs or min_chars <= 0:
            return paragraphs
        merged = []
        i = 0
        while i < len(paragraphs):
            cur = paragraphs[i].strip()
            if len(cur) < min_chars and i + 1 < len(paragraphs):
                # merge with next
                nxt = paragraphs[i + 1].strip()
                combined = (cur + " " + nxt).strip()
                merged.append(combined)
                i += 2
            else:
                merged.append(cur)
                i += 1
        return merged

    def split_into_paragraphs(text: str) -> list:
        """
        Convert raw section text into a list of paragraphs.
        Strategy:
          - First split on blank lines (existing paragraph separators).
          - For each resulting block:
            - If block is short, keep as-is.
            - If block is very long, perform sentence-level splitting and regroup into paragraphs of reasonable size.
          - Preserve bullet-like lists (lines starting with -, *, •).
        """
        if not text:
            return []
        t = text.replace("\r\n", "\n").replace("\r", "\n").strip()
        raw_blocks = [b.strip() for b in re.split(r"\n\s*\n", t) if b and b.strip()]
        paras: List[str] = []
        for block in raw_blocks:
            # detect bullet lists: many short lines starting with bullet chars
            lines = [ln.rstrip() for ln in block.split("\n") if ln.strip()]
            bullet_like = sum(1 for l in lines if re.match(r"^\s*[-\*\u2022]\s+", l))
            if bullet_like and bullet_like == len(lines):
                # convert each bullet line into its own paragraph with a bullet
                for l in lines:
                    cleaned = re.sub(r"^\s*[-\*\u2022]\s+", "", l).strip()
                    if cleaned:
                        paras.append(f"• {cleaned}")
                continue

            # if block is short enough, keep it
            if len(block) <= 520:
                paras.append(block)
                continue

            # otherwise split into sentences and regroup
            sentences = _sentence_split(block)
            if not sentences:
                # fallback split on fixed widths
                approx_chunks = [block[i : i + 480].strip() for i in range(0, len(block), 480)]
                paras.extend(approx_chunks)
                continue
            chunked = _chunk_sentences_into_paragraphs(sentences, max_chars=480)
            paras.extend(chunked)

        # Merge very short paragraphs to avoid fragmented text
        paras = _merge_short_paragraphs(paras, min_chars=MIN_PARAGRAPH_CHARS)
        return paras

    # ---------------- Images normalization & filtering ----------------
    raw_images = images if images is not None else sections.get("images") or []
    normalized: List[Dict[str, Any]] = []
    seen_paths = set()
    for im in raw_images:
        if not isinstance(im, dict):
            continue
        path = im.get("path") or im.get("file") or im.get("filepath")
        if not path:
            continue
        path = os.path.abspath(path)
        if path in seen_paths:
            continue  # dedupe exact same file
        if not os.path.exists(path):
            continue
        try:
            sz = os.path.getsize(path)
        except Exception:
            sz = 0
        # filter out trivially small images (likely thumbnails/placeholders)
        if sz < MIN_IMAGE_BYTES:
            continue
        seen_paths.add(path)
        normalized.append(
            {
                "path": path,
                "mime": im.get("mime"),
                "cell_index": im.get("cell_index"),
                "caption": im.get("caption"),
                "caption_hint": im.get("caption_hint"),
                "score": float(im.get("score")) if im.get("score") is not None else None,
                "section": im.get("section"),  # optional inline placement (string like 'results')
                "size": sz,
            }
        )

    # If LLM produced figure captions (JSON string), parse them.
    figure_captions: List[str] = []
    try:
        fc_raw = sections.get("figure_captions") or sections.get("figureCaptions") or None
        if fc_raw:
            if isinstance(fc_raw, str):
                # fc_raw might be JSON array string or plain bracketed strings
                try:
                    parsed = json.loads(fc_raw)
                    if isinstance(parsed, list):
                        figure_captions = [str(x).strip() for x in parsed]
                except Exception:
                    # fallback: try to extract quoted strings
                    extracted = re.findall(r'"([^"]+)"', fc_raw) or re.findall(r"'([^']+)'", fc_raw)
                    figure_captions = [str(x).strip() for x in extracted]
            elif isinstance(fc_raw, list):
                figure_captions = [str(x).strip() for x in fc_raw]
    except Exception:
        figure_captions = []

    # ---------------- Selection & ranking ----------------
    # Rank by: caption presence (explicit) > score > file size
    def usefulness_key(im: Dict[str, Any]):
        cap_flag = 1 if im.get("caption") else 0
        score = im.get("score") if im.get("score") is not None else 0.0
        size = im.get("size") or 0
        return (cap_flag, score, size)

    normalized.sort(key=usefulness_key, reverse=True)

    # limit to DEFAULT_MAX_FIGURES
    selected = normalized[:DEFAULT_MAX_FIGURES] if normalized else []

    # attach captions from figure_captions to selected images (in order)
    for idx, im in enumerate(selected):
        if not im.get("caption"):
            if idx < len(figure_captions) and figure_captions[idx].strip():
                im["caption"] = figure_captions[idx].strip()
        # fallback to caption_hint
        if not im.get("caption") and im.get("caption_hint"):
            im["caption"] = str(im.get("caption_hint")).strip()

    # Map images intended for inline insertion into sections; rest go to appendix
    images_by_section = {sec: [] for sec in ordered_sections}
    appendix_images: List[Dict[str, Any]] = []
    for im in selected:
        target = im.get("section")
        if target and isinstance(target, str):
            target_norm = target.strip().lower().replace(" ", "_")
            if target_norm in images_by_section:
                images_by_section[target_norm].append(im)
                continue
        appendix_images.append(im)

    # Keep a counter for numbering figures across inline + appendix
    figure_counter = 0

    # ---------------- Write sections and inline images ----------------
    for key in ordered_sections:
        content = sections.get(key, "")
        has_content = bool(content and content.strip())
        has_inline_images = bool(images_by_section.get(key))
        if not has_content and not has_inline_images:
            continue

        # add heading
        doc.add_heading(pretty_heading(key), level=1)

        # SPECIAL CASE: do not split abstract into multiple paragraphs; keep as one block
        if key == "abstract" and has_content:
            # Collapse extra whitespace but keep as a single paragraph, justified.
            single = re.sub(r"\s+", " ", content.strip())
            p = doc.add_paragraph(single)
            try:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except Exception:
                pass
        else:
            # write content paragraphs (justified)
            if has_content:
                paras = split_into_paragraphs(content)
                for para_text in paras:
                    # create paragraph
                    p = doc.add_paragraph()
                    # cleaned para
                    if isinstance(para_text, str) and para_text.lstrip().startswith("• "):
                        # bullet-like handled as plain paragraph with bullet char
                        run = p.add_run(para_text.strip())
                    else:
                        cleaned_para = re.sub(r"\s+", " ", para_text).strip()
                        run = p.add_run(cleaned_para)
                    try:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    except Exception:
                        pass

        # insert inline images for this section, if any
        for im in images_by_section.get(key, []):
            figure_counter += 1
            _insert_image_with_caption(doc, im, max_image_width_inch, figure_number=figure_counter)

    # ---------------- Figures appendix for remaining images ----------------
    if appendix_images:
        doc.add_page_break()
        doc.add_heading("Figures", level=1)
        for im in appendix_images:
            figure_counter += 1
            _insert_image_with_caption(doc, im, max_image_width_inch, figure_number=figure_counter)

    # ---------------- References ----------------
    refs = sections.get("references", "")
    if refs is not None:
        refs = refs.strip()

    doc.add_page_break()
    doc.add_heading("References", level=1)

    if refs:
        refs_norm = refs.replace("\r\n", "\n").replace("\r", "\n").strip()
        refs_collapsed = re.sub(r"\s+", " ", refs_norm).strip()
        # insert newline before numeric bracketed indices [1], [2] etc.
        refs_with_breaks = re.sub(r"\s*\[(\d+)\]\s*", r"\n[\1] ", refs_collapsed)
        refs_with_breaks = refs_with_breaks.lstrip("\n").strip()
        ref_lines = [line.strip() for line in refs_with_breaks.split("\n") if line and line.strip()]
        for line in ref_lines:
            p = doc.add_paragraph(line)
            try:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            except Exception:
                pass
    else:
        p = doc.add_paragraph("No references available.")
        try:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        except Exception:
            pass

    # ---------------- Save the document ----------------
    try:
        if isinstance(out_file, str):
            parent = os.path.dirname(out_file)
            if parent and not os.path.exists(parent):
                os.makedirs(parent, exist_ok=True)
            doc.save(out_file)
        elif isinstance(out_file, io.BytesIO):
            doc.save(out_file)
            out_file.seek(0)
        else:
            raise TypeError("out_file must be a str (path) or io.BytesIO")
    except Exception as e:
        raise RuntimeError(f"Error while saving DOCX: {e}")


# ---------------- Helper(s) ----------------
def _insert_image_with_caption(
    doc: Document, im: Dict[str, Any], max_image_width_inch: float, figure_number: Optional[int] = None
):
    """
    Insert an image into the doc with a small italic caption underneath.
    - Centers the image and caption.
    - Accepts an optional figure_number to prepend ("Figure N: ...").
    """
    path = im.get("path")
    caption_text = im.get("caption") or im.get("caption_hint") or ""
    # Prepend figure number if provided
    if figure_number is not None:
        if caption_text:
            caption_text = f"Figure {figure_number}: {caption_text}"
        else:
            caption_text = f"Figure {figure_number}"

    # Insert the image
    try:
        if Image is not None:
            try:
                with Image.open(path) as pil_img:
                    width_px, height_px = pil_img.size
                    info_dpi = pil_img.info.get("dpi")
                    dpi = None
                    if isinstance(info_dpi, tuple) and len(info_dpi) >= 1:
                        dpi = info_dpi[0]
                    elif isinstance(info_dpi, (int, float)):
                        dpi = float(info_dpi)
                    if not dpi or dpi <= 0:
                        dpi = 96.0
                    width_in = width_px / dpi if dpi else max_image_width_inch
                    target_width = min(max_image_width_inch, width_in if width_in > 0 else max_image_width_inch)
                    # Insert the image centered
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    try:
                        run = pic_para.add_run()
                        run.add_picture(path, width=Inches(target_width))
                    except Exception:
                        # fallback: add picture without computed width
                        run = pic_para.add_run()
                        run.add_picture(path, width=Inches(max_image_width_inch))
            except Exception:
                # fallback: insert without sizing calculation
                pic_para = doc.add_paragraph()
                pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = pic_para.add_run()
                run.add_picture(path, width=Inches(max_image_width_inch))
        else:
            # PIL not available: insert with max width
            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = pic_para.add_run()
            run.add_picture(path, width=Inches(max_image_width_inch))
    except Exception as e:
        # If insertion fails, add a placeholder paragraph (centered)
        err_para = doc.add_paragraph(f"[Could not insert image: {os.path.basename(path)} — {e}]")
        err_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return

    # Add small styled caption (centered)
    if caption_text:
        try:
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = cap_para.add_run(caption_text)
            run.italic = True
            run.font.size = Pt(9)
        except Exception:
            # fallback plain paragraph
            try:
                para = doc.add_paragraph(caption_text)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                pass
